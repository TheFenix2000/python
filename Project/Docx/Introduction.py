from docx import Document
from docx.shared import Cm


#otwieranie dokumentu
document = Document()

#tytuły i paragrafy
document.add_heading('This is a title of a document')
document.add_heading('This is a subtitle of a document', level=2)
paragraph = document.add_paragraph('Im a paragraph 1')
prior_paragraph = paragraph.insert_paragraph_before("Im inserted befor paragraph 1")

#"przerwa na stronie" - 'Enter'
document.add_page_break()
document.add_paragraph('Im a paragraph after a page break')
document.add_page_break()

#tabelki
table = document.add_table(rows=2,cols=2,style='Table Grid')
cell = table.cell(0,1)#row and column start from 0
cell.text = 'parrot,possibly dead' #row and column
row = table.rows[1]#dostę do wysztkich komórek w danym wierszu o numerze [1]
row.cells[0].text = 'I like you'
row.cells[1].text = 'I think I like you'

#wyświetlanie wartości każdej komórki w tabeli
for row in table.rows:#dla każdego wiersza w tabeli wypisz wartość każdej z komorek w tym wierszu
    for cell in row.cells:
        print(cell.text)

#wyliczanie ilości wierszy i kolumn w tabeli
row_count = len(table.rows)
col_count = len(table.columns)
print('Number of rows: {} \nNumber of columns: {}'.format(row_count,col_count))

#wstawianie wierszy do tabeli
row = table.add_row()

#dodawanie obrazka o określonym rozmiarze
document.add_picture('image.jpg',width=Cm(5.0))

#nadawanie stylu paragrafowi
paragraph2 = document.add_paragraph('I have a style on me')
paragraph2.style = 'ListBullet'

#dodawanie do istniejących paragrafów
paragraph3 = document.add_paragraph('Whatd u say?')
paragraph3.add_run('\nYEET')

#styl tekstu
    #nadawianie pogrubienia
paragraph4 = document.add_paragraph('Whatd u say? ')
run = paragraph4.add_run('YEET ')
run.bold = True
paragraph4.add_run('That is what I lke')
    #nadawanie pochylenia inną metodą
paragraph5 = document.add_paragraph('This is italic -> ')
paragraph5.add_run('Wee').italic = True
    #pogrubienie i pochylenie inna metoda
paragraph6 = document.add_paragraph('Im normal ')
paragraph6.add_run('Im bold ').bold = True
paragraph6.add_run('Im normal too ')
paragraph6.add_run('And Im italic').italic = True

#czcionki
    #metoda 1
paragraph7 = document.add_paragraph('Im a normal text ')
paragraph7.add_run('Im a cool text', 'Emphasis')
    #metoda 2
paragraph8 = document.add_paragraph('Hi there ')
run2 = paragraph8.add_run('I have an emphasis on me')
run2.style = 'Emphasis'


document.save('Introduction_document.docx')
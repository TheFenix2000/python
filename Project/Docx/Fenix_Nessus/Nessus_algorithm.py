#!/usr/bin/python

from sys import argv
import xml.etree.ElementTree as ET
import csv
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

DEFAULT_EXCLUDED_FILENAME = "excludedIDs.csv"

class NessusFile:
  def __init__(self, tree):
    self._tree = tree

  def find_items(self):
    full_list = []

    for host in self._tree.findall('Report/ReportHost'):
      self._ipaddr = host.find("HostProperties/tag/[@name='host-ip']").text

      for item in host.findall('ReportItem'):
        if item.find('cve') == None:
          cve = 'None'
        else:
          cve = item.find('cve').text
        if item.find('see_also') == None:
          see_also = 'None'
        else:
          see_also = item.find('see_also').text
        risk_factor = item.find('risk_factor').text
        pluginID = item.get('pluginID')
        pluginName = item.get('pluginName')
        description = item.find('description').text
        solution = item.find('solution').text
        order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "None": 4}

        eight = (order[risk_factor], pluginName, risk_factor, cve, description, solution, see_also, pluginID)
        full_list.append(eight)

    full_list.sort()
    return full_list

  def get_ipaddr(self):
    return self._ipaddr

class ExcludedLoader:
  def __init__(self, filename=None):
    if filename is not None:
      self._filename = filename
    else:
      self._filename = DEFAULT_EXCLUDED_FILENAME

  def get_excluded(self):
    return self._get_excluded_ids_from_file(self._filename)

  def _get_excluded_ids_from_file(self, filename):
    with open(filename, 'r') as file:
      csv_file = csv.reader(file)
      for row in csv_file:
        dictionary = {i: True for i in row}
        return dictionary

class NessusProcess:
  def __init__(self, excluded_loader, ipaddr):
    self._excluded_loader = excluded_loader
    self._ipaddr = ipaddr

  def _shade_cells(self, cell, condition):
    if condition == 'Critical':
      shade = '#960000'
    elif condition == 'High':
      shade = '#fa3434'
    elif condition == 'Medium':
      shade = '#ffa929'
    else:
      shade = 'ffff40'
    tcPr = cell._tc.get_or_add_tcPr()
    tcVAlign = OxmlElement("w:shd")
    tcVAlign.set(qn("w:fill"), shade)
    tcPr.append(tcVAlign)

  def process(self, full_list):
    excluded = self._excluded_loader.get_excluded()
    ipaddr = self._ipaddr
    result_list = [ x for x in full_list if x[1] not in excluded and x[0] < 4 ]

    path = 'document.docx'
    document = Document(path)
    for i in range(len(result_list)):
      document.add_paragraph()
      paragraph1 = document.add_paragraph()
      run1 = paragraph1.add_run("1." + str(i+1) + ".     " + result_list[i][1])
      run1.bold = True
      run1.font.size = Pt(14)
      document.add_paragraph()

      table = document.add_table(cols=2, rows=2)
      table.alignment = WD_TABLE_ALIGNMENT.RIGHT
      for row in table.rows:
        for cell in row.cells:
          cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
      table.columns[0].width = Cm(7.0)
      table.columns[1].width = Cm(9.0)
      for row in table.rows:
        row.height = Cm(0.8)

      run = table.cell(0, 0).paragraphs[0].add_run('Rick Score:')
      run.bold = True
      run.font.size = Pt(12)
      table.cell(0, 1).paragraphs[0].add_run(result_list[i][2])
      self._shade_cells(table.cell(0, 1), table.cell(0, 1).text)

      run10 = table.cell(1, 0).paragraphs[0].add_run('Affected Systems:')
      table.cell(1, 1).paragraphs[0].add_run('https://' + ipaddr)
      run10.bold = True
      run10.font.size = Pt(12)

      if result_list[i][3] != 'None':
        table.add_row()
        run = table.cell(2, 0).paragraphs[0].add_run('CVSS Risk Score:')
        table.cell(2, 1).paragraphs[0].add_run(result_list[i][3])
        run.bold = True
        run.font.size = Pt(12)
        table.rows[2].height = Cm(0.8)
        table.rows[2].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        table.rows[2].cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
      document.add_paragraph()

      paragraph11 = document.add_paragraph()
      run11 = paragraph11.add_run("1." + str(i+1) + ".1.   " + "Vulnerability Description")
      run11.bold = True
      run11.font.size = Pt(13)
      document.add_paragraph()
      paragraph11 = document.add_paragraph(result_list[i][4])
      paragraph11.paragraph_format.left_indent = Cm(1.4)
      document.add_paragraph()

      paragraph12 = document.add_paragraph()
      run12 = paragraph12.add_run("1." + str(i + 1) + ".2.   " + "Evidence")
      run12.bold = True
      run12.font.size = Pt(13)
      document.add_paragraph()
#to be added
      document.add_paragraph()

      paragraph13 = document.add_paragraph()
      run13 = paragraph13.add_run("1." + str(i + 1) + ".3.   " + "Recommendation")
      run13.bold = True
      run13.font.size = Pt(13)
      document.add_paragraph()
      paragraph13 = document.add_paragraph(result_list[i][5])
      paragraph13.paragraph_format.left_indent = Cm(1.4)
      document.add_paragraph()

      if result_list[i][7] != 'None':
        paragraph14 = document.add_paragraph()
        run14 = paragraph14.add_run("1." + str(i + 1) + ".4.   " + "References")
        run14.bold = True
        run14.font.size = Pt(13)
        document.add_paragraph()
        paragraph14 = document.add_paragraph(result_list[i][6])
        paragraph14.paragraph_format.left_indent = Cm(1.4)
      document.add_page_break()

    document.save('Nessus-result.docx')

def main(nessus_filename):
  """Main function"""
  tree = ET.parse(nessus_filename)

  nessus = NessusFile(tree)
  full_list = nessus.find_items()
  ipaddr = nessus.get_ipaddr()
  ld = ExcludedLoader()
  n = NessusProcess(ld, ipaddr)
  n.process(full_list)


if __name__ == '__main__':
  nessus_filename = argv[1]
  main(nessus_filename)
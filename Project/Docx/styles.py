from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm,Pt


document = Document()

styles = document.styles
paragraph_stles = [
    s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH
]                                                          #wyświetlanie stylów wbudowanych
#for style in paragraph_stles:
    #print(style.name)

#używanie styli
paragraph = document.add_paragraph('Hi there')
paragraph.style = document.styles['Heading 1']

#zdefiniuj formatowanie paragrafu
style = document.styles.add_style('Indent', WD_STYLE_TYPE.PARAGRAPH)
paragraph_format = style.paragraph_format
paragraph_format.left_indent = Cm(0.5)
paragraph_format.first_line_indent = Cm(-0.5)
paragraph_format.space_before = Pt(18)
paragraph_format.widow_control = True
paragraph2 = document.add_paragraph('Hello, its me Mario')
paragraph2.style = 'Indent'
paragraph3 = document.add_paragraph('im normal text')



document.save('style.docx')
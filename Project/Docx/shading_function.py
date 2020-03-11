from docx import Document
from docx.oxml.shared import OxmlElement, qn
from docx.shared import RGBColor
def shade_cells(cells, shade):
    for cell in cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcVAlign = OxmlElement("w:shd")
        tcVAlign.set(qn("w:fill"), shade)
        tcPr.append(tcVAlign)


document = Document()
table = document.add_table(rows=3, cols=3)

shade_cells([table.cell(0, 0), table.cell(1, 1), table.cell(2, 2)], "#99badd")  # unc
shade_cells([table.cell(0, 1), table.cell(1, 2)], "#CC0000")  # nc state
shade_cells([table.cell(1, 0), table.cell(2, 1)], "#fcba03")  # duke

run = table.cell(0, 0).paragraphs[0].add_run('axd')
font_c = run.font.color
font_c.rgb = RGBColor(221, 32, 7)

document.save('test4.docx')
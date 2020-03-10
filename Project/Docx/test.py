from docx import Document
from docx.shared import Pt
import os
list_of_elements = [("4", "Hi", "Hello", "Nice to meet you"), ("3", "Not Hi", "Terrible", "I'm sick of you")]

path = os.path.join('document.docx')
doc = Document(path)
doc.paragraphs[0].add_run('Hi there')

for i in range(len(list_of_elements)):
    paragraph1 = doc.add_paragraph()
    run = paragraph1.add_run(str(i+1) + ". " + list_of_elements[i][1])
    run.bold = True
    run.font.size = Pt(10)
    paragraph = doc.add_paragraph()
    paragraph.add_run(list_of_elements[i][2])
doc.save('test2.docx')
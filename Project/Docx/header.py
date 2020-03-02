from docx import Document
from docx.shared import Inches, Pt

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

document = Document()
header = document.sections[0].header
p1 = header.add_paragraph()
p2 = header.add_paragraph()
r = p1.add_run()
r2 = p2.add_run()
r.add_picture('crossword.png', width=Inches(0.6))
r2.add_text('_________________________________________________________________________________________________________')
format = p2.paragraph_format
format.line_spacing = Pt(3)
p3 = header.paragraphs[0]
delete_paragraph(p3)

document.save('header.docx')
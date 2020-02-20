from docx import Document
from docx.shared import Cm,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_UNDERLINE
from docx.shared import RGBColor


document = Document()

#wyrównywanie tekstu w poziomie
paragraph = document.add_paragraph('Hi, Im a paragraph')
paragraph2 = document.add_paragraph('Hi, Im a paragraph the II')
paragraph_format = paragraph.paragraph_format #definiujemy mienną paragraph_format i przypisujemy do naszego paragrafu
paragraph2_format = paragraph2.paragraph_format
# indicating alignment is inherited from the style hierarchy - "None"
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph2_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

#wcięcia
paragraph_normal = document.add_paragraph('Normal text')
paragraph3 = document.add_paragraph('Nice to meet u, I have a indent of 0.5cm')
paragraph3_format = paragraph3.paragraph_format
paragraph3_format.left_indent = Cm(0.5)
document.add_page_break()

#nadawanie czcionki
run = document.add_paragraph('It is 14:12 ').add_run('That is Nice')
font = run.font
font.name = 'Calibri'
font.size = Pt(18)

#podkreślenie
run2 = document.add_paragraph('Im not underlined ').add_run('Im undelined but value NONE')
font2 = run2.font
font2.underline = WD_UNDERLINE.NONE

#kolorowu tekst
run3 = document.add_paragraph('Im black ').add_run('I have a color on me')
font3 = run3.font
font3.color.rgb = RGBColor(0x42,0x24,0xE9)



document.save('text.docx')
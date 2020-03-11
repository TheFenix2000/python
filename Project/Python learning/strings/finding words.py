from docx import Document
from docx.shared import Pt

lis = [('great', 'Its a beautiful day', 'crazy night'), ('Pog', 'Oww man its creepper', 'peaceful day')]

def create_paragraph(document, list_of_elements, font_weight, i, position):
    paragraph = document.add_paragraph()
    run = paragraph.add_run('1. ' + str(i+1) + ".     " + list_of_elements[i][position])
    run.bold = font_weight
    run.font.size = Pt(14)
    document.add_paragraph()

document = Document()
for i in range(len(lis)):
    create_paragraph(document, lis, True, i, 1)
document.save('demo.docx')